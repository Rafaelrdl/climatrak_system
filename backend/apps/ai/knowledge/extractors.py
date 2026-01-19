"""
Text Extractors for AI Knowledge Base (AI-006).

Extrai texto de diferentes formatos de arquivo:
- PDF: via pypdf
- DOCX: via python-docx
- Markdown: texto puro

Todos os extractors seguem a interface:
    extract(file_obj) -> str
"""

import io
import logging
import re
from abc import ABC, abstractmethod
from typing import BinaryIO, Union

logger = logging.getLogger(__name__)


class BaseExtractor(ABC):
    """Interface base para extractors de texto."""

    @abstractmethod
    def extract(self, file_obj: BinaryIO) -> str:
        """
        Extrai texto do arquivo.

        Args:
            file_obj: Objeto de arquivo binário (file-like object)

        Returns:
            Texto extraído, limpo e normalizado
        """
        pass

    def normalize_text(self, text: str) -> str:
        """
        Normaliza texto extraído.

        - Remove caracteres de controle
        - Normaliza espaços e quebras de linha
        - Remove linhas em branco excessivas
        """
        if not text:
            return ""

        # Remove caracteres de controle (exceto newline e tab)
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]", "", text)

        # Normaliza quebras de linha
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        # Remove espaços no final das linhas
        text = re.sub(r"[ \t]+$", "", text, flags=re.MULTILINE)

        # Limita linhas em branco consecutivas a 2
        text = re.sub(r"\n{3,}", "\n\n", text)

        # Remove espaços no início e fim
        text = text.strip()

        return text


class PDFExtractor(BaseExtractor):
    """
    Extrator de texto para arquivos PDF.

    Usa pypdf para extrair texto de todas as páginas.
    """

    def extract(self, file_obj: BinaryIO) -> str:
        """
        Extrai texto de PDF.

        Args:
            file_obj: Arquivo PDF em modo binário

        Returns:
            Texto extraído de todas as páginas
        """
        try:
            from pypdf import PdfReader

            # Garante que estamos no início do arquivo
            file_obj.seek(0)

            reader = PdfReader(file_obj)
            text_parts = []

            for page_num, page in enumerate(reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num}: {e}")
                    continue

            full_text = "\n\n".join(text_parts)
            return self.normalize_text(full_text)

        except ImportError:
            raise ImportError("pypdf is required for PDF extraction. Install with: pip install pypdf")
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise ExtractionError(f"Failed to extract PDF: {e}") from e


class DocxExtractor(BaseExtractor):
    """
    Extrator de texto para arquivos DOCX (Word).

    Usa python-docx para extrair texto de parágrafos e tabelas.
    """

    def extract(self, file_obj: BinaryIO) -> str:
        """
        Extrai texto de DOCX.

        Args:
            file_obj: Arquivo DOCX em modo binário

        Returns:
            Texto extraído de parágrafos e tabelas
        """
        try:
            from docx import Document

            # Garante que estamos no início do arquivo
            file_obj.seek(0)

            doc = Document(file_obj)
            text_parts = []

            # Extrai parágrafos
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extrai texto de tabelas
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            full_text = "\n\n".join(text_parts)
            return self.normalize_text(full_text)

        except ImportError:
            raise ImportError("python-docx is required for DOCX extraction. Install with: pip install python-docx")
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise ExtractionError(f"Failed to extract DOCX: {e}") from e


class MarkdownExtractor(BaseExtractor):
    """
    Extrator de texto para arquivos Markdown.

    Como Markdown já é texto, apenas normaliza.
    Opcionalmente remove sintaxe Markdown.
    """

    def __init__(self, strip_markdown: bool = False):
        """
        Args:
            strip_markdown: Se True, remove sintaxe Markdown (headers, links, etc)
        """
        self.strip_markdown = strip_markdown

    def extract(self, file_obj: BinaryIO) -> str:
        """
        Extrai texto de Markdown.

        Args:
            file_obj: Arquivo Markdown em modo binário

        Returns:
            Texto (com ou sem sintaxe Markdown)
        """
        try:
            # Garante que estamos no início do arquivo
            file_obj.seek(0)

            # Tenta diferentes encodings
            content = None
            for encoding in ["utf-8", "latin-1", "cp1252"]:
                try:
                    file_obj.seek(0)
                    content = file_obj.read().decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if content is None:
                raise ExtractionError("Failed to decode Markdown file")

            if self.strip_markdown:
                content = self._strip_markdown_syntax(content)

            return self.normalize_text(content)

        except Exception as e:
            if isinstance(e, ExtractionError):
                raise
            logger.error(f"Markdown extraction failed: {e}")
            raise ExtractionError(f"Failed to extract Markdown: {e}") from e

    def _strip_markdown_syntax(self, text: str) -> str:
        """
        Remove sintaxe Markdown comum.

        Mantém o texto legível, removendo:
        - Headers (#)
        - Bold/italic (**/*/_)
        - Links [text](url)
        - Images ![alt](url)
        - Code blocks ```
        - Inline code `
        """
        # Remove code blocks
        text = re.sub(r"```[\s\S]*?```", "", text)

        # Remove inline code (mantém conteúdo)
        text = re.sub(r"`([^`]+)`", r"\1", text)

        # Remove images
        text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)

        # Remove links (mantém texto)
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)

        # Remove headers (mantém texto)
        text = re.sub(r"^#+\s*", "", text, flags=re.MULTILINE)

        # Remove bold/italic
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"\*([^*]+)\*", r"\1", text)
        text = re.sub(r"__([^_]+)__", r"\1", text)
        text = re.sub(r"_([^_]+)_", r"\1", text)

        # Remove horizontal rules
        text = re.sub(r"^[\-*_]{3,}\s*$", "", text, flags=re.MULTILINE)

        # Remove blockquotes
        text = re.sub(r"^>\s*", "", text, flags=re.MULTILINE)

        return text


class ExtractionError(Exception):
    """Erro durante extração de texto."""

    pass


# Factory function para escolher extractor baseado no tipo
EXTRACTORS = {
    "PDF": PDFExtractor,
    "DOCX": DocxExtractor,
    "MARKDOWN": MarkdownExtractor,
}


def extract_text(file_obj: BinaryIO, file_type: str) -> str:
    """
    Extrai texto de arquivo baseado no tipo.

    Args:
        file_obj: Objeto de arquivo binário
        file_type: Tipo do arquivo (PDF, DOCX, MARKDOWN)

    Returns:
        Texto extraído e normalizado

    Raises:
        ValueError: Se tipo de arquivo não suportado
        ExtractionError: Se falhar na extração
    """
    file_type = file_type.upper()

    if file_type not in EXTRACTORS:
        raise ValueError(f"Unsupported file type: {file_type}. Supported: {list(EXTRACTORS.keys())}")

    extractor_class = EXTRACTORS[file_type]
    extractor = extractor_class()

    return extractor.extract(file_obj)


def extract_text_from_path(file_path: str, file_type: str = None) -> str:
    """
    Extrai texto de arquivo por caminho.

    Args:
        file_path: Caminho do arquivo
        file_type: Tipo do arquivo (auto-detectado se não fornecido)

    Returns:
        Texto extraído e normalizado
    """
    import os

    if file_type is None:
        ext = os.path.splitext(file_path)[1].lower()
        type_map = {
            ".pdf": "PDF",
            ".docx": "DOCX",
            ".md": "MARKDOWN",
            ".markdown": "MARKDOWN",
        }
        file_type = type_map.get(ext)
        if file_type is None:
            raise ValueError(f"Cannot auto-detect type for extension: {ext}")

    with open(file_path, "rb") as f:
        return extract_text(f, file_type)
